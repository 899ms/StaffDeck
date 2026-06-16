from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/enterprise_service_knowledge_base_longform.docx")


def set_east_asia_font(style, font_name: str) -> None:
    style.font.name = font_name
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:eastAsia"), font_name)


def add_paragraph(doc: Document, text: str, style: str = "Normal") -> None:
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(level=level)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(46, 116, 181)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(46, 116, 181)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(31, 77, 120)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
        style = doc.styles[style_name]
        set_east_asia_font(style, "PingFang SC")
        if style_name == "Normal":
            style.font.name = "Arial"
            style.font.size = Pt(11)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.line_spacing = 1.25


def chapter_paragraphs(chapter: str, domain: str, concern: str, examples: list[str]) -> list[str]:
    example_text = "、".join(examples)
    return [
        (
            f"{chapter}的知识用于帮助模型在没有一次性读取全文的情况下逐步定位答案。"
            f"当用户围绕{domain}提出问题时，系统应先判断问题属于政策解释、状态核对、异常排查还是执行动作，"
            f"再展开相应知识桶。处理时不要只看单句关键词，也不要把示例当成固定脚本；应结合用户身份、历史会话、"
            f"订单状态、当前智能体可见资源和业务上下文，形成可以追溯的判断。常见触发包括{example_text}。"
        ),
        (
            f"在{domain}场景下，客服回复必须同时满足准确、可执行和边界清楚三项要求。"
            f"如果文档只给出原则而没有执行条件，应先说明已知信息，再指出还需要补充哪些字段；"
            f"如果文档给出了明确条件，应基于条件给出结论，不要把“建议稍后查看”当成最终答案。"
            f"涉及{concern}时，需要优先保护用户体验，同时保留系统审计需要的字段。"
        ),
        (
            f"渐进式知识检索时，第一轮应只读取章节摘要和桶摘要；第二轮再读取最相关的片段；"
            f"如果片段之间存在冲突，优先采用发布时间更新、适用范围更窄、与用户条件更匹配的内容。"
            f"当需要调用工具或转人工时，应把文档片段作为依据交给执行模型，而不是直接拼接原文。"
            f"{chapter}中的问答样例只用于辅助识别，不代表最终回复模板。"
        ),
    ]


CHAPTERS = [
    (
        "第一章 企业服务定位与知识使用边界",
        "企业服务定位",
        "品牌承诺、服务语气和可解释性",
        ["用户问公司能做什么", "用户质疑客服身份", "用户要求解释处理依据"],
        [
            "面壁智能客服面向企业级客户、普通消费者和内部运营同事提供统一服务入口。它不是单纯问答机器人，而是根据智能体可见域选择技能、工具、知识库和通用能力的对话运行时。",
            "所有知识回答都应区分事实、政策、建议和执行结果。事实来自知识库、工具或会话记忆；政策来自已发布文档；建议用于引导下一步；执行结果必须来自工具或人工处理记录。",
            "当用户问题跨越多个域，例如同时问会员权益、配送承诺和退款进度，模型应先拆成可管理的任务，再按优先级推进。不能为了让回复显得完整而编造未检索到的细节。",
        ],
    ),
    (
        "第二章 用户身份、会话记忆与称呼规则",
        "身份识别与记忆",
        "姓名、会员等级、联系方式和授权边界",
        ["用户说我叫某某", "用户更改称呼", "用户要求使用历史地址"],
        [
            "长期记忆只记录稳定、可复用且用户不反感的信息，例如称呼偏好、常用会员账号、默认收货城市和沟通偏好。一次性订单号、临时投诉情绪和单次活动状态不应进入长期记忆。",
            "同一用户多次更新称呼时，应保留最新称呼，并在回复中自然使用。若记忆与当前消息冲突，当前消息优先；若用户明确要求不要记住，应停止写入。",
            "涉及账号、手机号、地址等敏感字段时，模型可以提示需要用户确认，但不应在公开回复中重复完整敏感信息。必要时仅展示脱敏片段。",
        ],
    ),
    (
        "第三章 商品、价格与比价咨询",
        "商品和价格",
        "价格来源、库存状态和比价口径",
        ["A1 和 A3 哪个便宜", "iPhone 15 多少钱", "商品为什么涨价"],
        [
            "商品价格必须以工具结果、已发布价目表或知识库价格规则为依据。没有依据时应说明需要查询，而不是直接给出价格。",
            "比价时要说明比较口径，包括商品名称、规格、版本、币种、是否含税、是否包含促销券和时间范围。不同渠道价格不可直接混用。",
            "若用户同时表达购买和比价，系统应先明确用户的主目标：立即购买、先比较后购买，或为已购商品询价。必要时可拆分成购买任务和比价任务，但不能重复创建同一个任务。",
        ],
    ),
    (
        "第四章 订单创建、支付与取消",
        "订单履约",
        "下单确认、金额核验和取消窗口",
        ["帮我买一个 A1", "订单已支付了吗", "我想取消刚才订单"],
        [
            "创建订单前必须确认关键字段：用户身份、商品标识、数量、收货或交付方式，以及是否同意下单。若用户一句话已提供多个字段，应一次性落槽，不要重复追问。",
            "支付状态分为未支付、已支付、支付处理中、支付失败和退款中。系统展示金额时必须来自订单工具或已确认订单记录，不能用用户口述金额覆盖工具返回。",
            "取消订单要看发货状态和支付状态。未支付订单通常可直接取消；已支付未发货订单需要走取消和退款流程；已发货订单通常进入退货或拒收流程。",
        ],
    ),
    (
        "第五章 物流履约、配送承诺与改派",
        "物流配送",
        "时效承诺、地址变更和仓配评估",
        ["为什么还没发货", "能不能改地址", "承诺今天到为什么没到"],
        [
            "配送问题要先区分未发货、已出库、运输中、派送中、已签收和异常签收。不同状态下可执行动作不同，不能把查询物流和修改地址混为同一步。",
            "地址变更需要确认订单号、用户身份、目标地址、期望送达时间和是否接受延迟。若仓配评估显示不可改派，应给出原因和替代方案。",
            "承诺时效只在活动页、订单页或客服工具返回中明确存在时生效。用户截图可作为线索，但仍需要工具核对订单承诺字段。",
        ],
    ),
    (
        "第六章 退款、退货与换货",
        "售后处理",
        "权益条件、商品状态和用户确认",
        ["我要退款", "我要退货", "换一个颜色", "东西坏了"],
        [
            "退款、退货和换货不是同一个动作。退款强调资金返还，退货强调商品回收，换货强调重新履约。用户只说“退一下”时，需要识别其真实诉求或追问。",
            "售后处理前需要订单号、商品状态、签收时间、原因、凭证和用户期望。若订单状态已经满足条件，可以继续处理；若信息不足，应明确缺失字段。",
            "涉及退款金额、退货运费、换货库存时，必须查询工具或知识片段。不能承诺“一定成功”，只能说明当前条件下可申请、可尝试或需人工审核。",
        ],
    ),
    (
        "第七章 会员等级、权益发放与补偿",
        "会员权益",
        "应发权益、实发权益和差异补偿",
        ["会员券没到账", "积分少了", "活动礼品没发"],
        [
            "会员权益核对要明确用户 ID、订单号、会员等级、权益类型和活动批次。若用户只说“权益少了”，需要先确认权益类型和发生场景。",
            "应发权益和实发权益不一致时，应说明差异来源，例如活动资格不满足、发放延迟、渠道限制、库存不足或系统任务失败。能够自动补发时再进入补偿动作。",
            "补偿不是所有问题的默认结果。只有文档或工具返回允许自动补偿，且用户身份和订单条件满足时，才能执行；否则应转人工复核或告知等待任务完成。",
        ],
    ),
    (
        "第八章 投诉、风险与人工介入",
        "投诉风险",
        "升级路径、证据保全和响应时限",
        ["我要投诉", "我要找人工", "你们欺骗消费者"],
        [
            "投诉处理要先安抚情绪，再收集问题事实。应记录用户诉求、涉及订单、证据、期望结果和是否存在安全风险。",
            "高风险内容包括人身安全、法律威胁、媒体曝光、重复扣款、未授权交易、隐私泄露和大额损失。出现这些内容时应提高优先级，并保留完整事件摘要。",
            "转人工不是失败，而是明确责任边界。模型应说明转人工原因、已收集信息和建议处理方向，避免让用户重复描述。",
        ],
    ),
    (
        "第九章 企业文化、品牌语气与服务红线",
        "企业文化",
        "语气一致性、承诺边界和禁用表达",
        ["客服态度不好", "能不能说人话", "给我一个保证"],
        [
            "企业服务语气要求清楚、负责、不过度承诺。可以表达理解和歉意，但不能承认未核实的事实，也不能暗示用户必须接受某个方案。",
            "禁用表达包括推卸责任、威胁用户、夸大处理权限、承诺必然成功、泄露内部评分规则和让用户自行承担系统错误。",
            "对外回复应优先给出下一步，而不是解释内部流程。只有用户要求了解依据时，才展开政策条款或知识来源。",
        ],
    ),
    (
        "第十章 内部运营排查与数据口径",
        "运营排查",
        "指标口径、工单归因和复盘字段",
        ["为什么差评多", "这个问题归谁", "怎么复盘"],
        [
            "运营复盘关注问题来源、影响范围、触发路径、处理耗时、用户损失和改进动作。单次对话的点赞点踩不能直接代表技能质量，需要结合模型判别。",
            "问题分桶可以分为模型理解问题、技能设计问题、工具或系统问题、用户随意反馈和已解决反馈。模型应根据上下文判断，不要只看点踩按钮。",
            "统计指标应明确分母。差评率以调用次数或有效会话数为分母时含义不同；一个流程内多次点踩通常应按一次会话级负反馈计算。",
        ],
    ),
    (
        "第十一章 知识检索、工具发现与技能生成",
        "知识工程",
        "分桶切片、工具建议和技能草稿",
        ["上传文档怎么变技能", "工具怎么发现", "为什么知识没命中"],
        [
            "知识库文档应先按章节、主题和段落切成桶，再把桶拆成片段。模型先读桶摘要，再按需读取片段，避免长文档一次性进入上下文。",
            "文档中若自然描述了可执行接口，例如给出请求地址、方法、输入字段和返回字段，系统可以生成工具草案。未确认前不得写入技能图的调用动作。",
            "场景化技能从知识中发现后必须经过用户确认、改写和发布。发现结果只是建议，不代表运行时已经具备执行能力。",
        ],
    ),
    (
        "第十二章 QA 验收、边界样例与回归问题",
        "质量验收",
        "问答覆盖、边界条件和回归检查",
        ["怎么验证知识库", "为什么回答旧答案", "如何检查切片"],
        [
            "知识库验收应覆盖精确事实问答、跨章节综合问答、缺失信息追问、冲突知识裁决和工具建议发现。每类问题至少准备三到五个样例。",
            "如果用户问题命中文档但模型没有展开相关桶，应检查桶摘要是否过于笼统、问题是否需要语义改写、以及智能体是否绑定了正确知识库版本。",
            "回归检查要确认下线的知识库仍可在管理页看到，分支智能体只影响自己的可见域，整体智能体删除才是真正全局删除。",
        ],
    ),
]


QA_ITEMS = [
    ("用户问“我叫 hm，以后叫我老师”，应如何处理？", "更新长期称呼记忆为老师，当前轮可自然使用，不要把一次性订单信息写入长期记忆。"),
    ("用户问“A1 和 A3 哪个便宜”，可以直接回答价格吗？", "不可以。必须基于价格工具、价目表或知识片段查询后回答，并说明比较口径。"),
    ("用户说“我想退款，买贵了”，应先做什么？", "识别为售后退款诉求，收集或确认订单号、商品、签收状态和退款原因，不能承诺一定成功。"),
    ("用户同时说“退货吧，买完再买 A1”，任务如何处理？", "拆成售后任务和购买任务，先处理当前确认度更高或风险更高的任务，另一个进入任务队列，避免重复创建。"),
    ("会员券未到账需要哪些字段？", "用户 ID、订单号、会员等级、权益类型和活动批次；缺失时追问。"),
    ("配送承诺没有履约时可以直接补偿吗？", "不能。需要核对承诺来源、订单状态、仓配原因和补偿规则，工具或政策允许后才执行。"),
    ("知识库文档中出现接口地址但工具未配置怎么办？", "生成工具建议并等待用户确认；确认前不得把调用动作写入技能图。"),
    ("用户点踩一次是否等于技能差评？", "不一定。需要异步分析反馈归因，区分模型、技能、工具、用户随意反馈和已解决反馈。"),
    ("旧版 .doc 上传失败怎么办？", "提示转换为 .docx；当前解析链路支持 docx、txt、md、html、pdf 等，旧 doc 需要转换。"),
    ("知识冲突时如何裁决？", "优先更新时间更新、适用范围更窄、与用户条件更匹配的内容，并在必要时说明依据。"),
]


def build_document() -> Document:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title_run = title.add_run("面壁智能客户服务知识库长文测试手册")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("用于验证知识库上传、分桶、切片、渐进检索与 QA 问答能力")
    subtitle_run.font.name = "Arial"
    subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(85, 85, 85)

    add_paragraph(
        doc,
        "版本：2026-06-16 测试版。本文档故意覆盖企业介绍、商品、订单、物流、售后、会员、投诉、知识工程和运营复盘等多个主题，"
        "用于制造足够长、章节清晰但语义交叉的知识输入。模型应通过分桶摘要逐步发现相关片段，而不是一次性读取全文。",
    )
    add_heading(doc, "使用说明", 1)
    for item in [
        "适合作为企业端知识库上传测试文件，重点观察文档解析、章节分桶、段落切片、桶摘要和 QA 召回效果。",
        "文档包含多个相近概念，例如退款、退货、换货、取消订单和补偿，便于测试模型是否能做语义区分。",
        "文末提供 QA 样例，但样例不是固定话术；回答时仍应基于检索到的知识片段组织自然语言。",
    ]:
        add_bullet(doc, item)

    for chapter_index, (chapter, domain, concern, examples, anchors) in enumerate(CHAPTERS, start=1):
        add_heading(doc, chapter, 1)
        for anchor in anchors:
            add_paragraph(doc, anchor)
        for section_index in range(1, 5):
            add_heading(doc, f"{chapter_index}.{section_index} {domain}处理要点 {section_index}", 2)
            paragraphs = chapter_paragraphs(chapter, domain, concern, examples)
            for paragraph_index, paragraph in enumerate(paragraphs, start=1):
                suffix = (
                    f"本段用于测试第 {chapter_index} 章第 {section_index} 节第 {paragraph_index} 段的召回稳定性。"
                    f"如果用户问题只命中其中一个条件，模型应只展开必要片段；如果问题同时涉及{domain}和其他章节，"
                    "应在检索后合并证据，并说明哪些内容来自知识库、哪些内容还需要工具或用户补充。"
                )
                add_paragraph(doc, paragraph + suffix)
            add_heading(doc, f"{chapter_index}.{section_index}.1 可检索字段", 3)
            field_texts = [
                f"主题字段：{domain}、{concern}、{examples[0]}。",
                f"上下文字段：用户身份、会话记忆、智能体可见域、当前任务状态、工具返回结果。",
                f"边界字段：是否需要人工、是否需要补充信息、是否允许自动执行、是否存在风险升级。",
            ]
            for item in field_texts:
                add_bullet(doc, item)

    add_heading(doc, "专题 QA 锚点：用于检索验收的自然语言问题", 1)
    add_paragraph(
        doc,
        "本章把若干高频测试问题写成普通正文段落，用于验证长文档分桶后是否能从自然语言问题召回对应知识。"
        "这些内容不是回复模板，而是提供明确的政策、字段和处理边界。系统回答时仍应结合当前用户消息、会话状态、"
        "智能体可见域、工具结果和已召回的其他知识片段进行组织。",
    )
    for question, answer in QA_ITEMS:
        add_heading(doc, f"QA：{question}", 2)
        add_paragraph(
            doc,
            f"当用户提出“{question}”或语义相近的问题时，知识库应召回本段作为依据。处理要点：{answer}"
            "如果当前用户没有提供足够字段，应先追问缺失信息；如果已有工具结果或会话证据，应结合证据给出更具体的下一步。"
            "模型不得把这段话逐字返回给用户，而应转换为自然、简洁、可执行的客服回复。",
        )

    add_heading(doc, "附录 A：问答样例索引", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "问题"
    header[1].text = "期望答案要点"
    for question, answer in QA_ITEMS:
        cells = table.add_row().cells
        cells[0].text = question
        cells[1].text = answer

    add_heading(doc, "附录 B：用于长文档分桶的补充段落", 1)
    supplement_topics = [
        "跨域问题通常不是单条规则能解决。例如用户说“会员礼品没到，订单又晚了，还想取消”，模型需要先拆分会员权益、物流承诺和订单取消三个维度，再判断是否有共享字段。",
        "大文档切片时，段落边界比固定字数更重要。固定字数可以控制上下文预算，但不能破坏一条规则的完整条件、例外和处理动作。",
        "知识检索结果进入回复前，应经过模型整合。系统不应该把片段原文直接贴给用户，除非用户明确要求查看依据。",
        "当知识库版本发生变化时，旧会话可能仍引用旧知识。测试时应区分知识版本、智能体分支和当前会话绑定的智能体。",
    ]
    for repeat in range(18):
        for topic in supplement_topics:
            add_paragraph(
                doc,
                f"补充段落 {repeat + 1}：{topic}"
                "这段内容用于增加文档长度和主题重叠度，便于测试检索系统在相似表述中选择正确知识桶。"
                "如果 QA 问题只需要其中一句，应返回最小充分证据；如果需要综合多段，应在最终回答里保持结论明确。",
            )

    add_heading(doc, "附录 C：快速 QA 锚点", 1)
    add_paragraph(
        doc,
        "本附录故意放在文档末尾，用于验证无模型配置时最近知识桶 fallback 也能读取到可问答片段。"
        "这些锚点重复正文中的高频问题，便于观察知识库上传后的最小 QA 可用性。",
    )
    for question, answer in QA_ITEMS:
        add_paragraph(
            doc,
            f"快速问答锚点：{question} 参考要点：{answer}"
            "回答时应先给出结论，再补充需要的字段、证据来源或下一步动作；不要直接照抄文档句子。",
        )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    cjk_count = len(re.findall(r"[\u4e00-\u9fff]", text))
    print(f"written={OUTPUT}")
    print(f"paragraphs={len(doc.paragraphs)}")
    print(f"tables={len(doc.tables)}")
    print(f"cjk_chars={cjk_count}")


if __name__ == "__main__":
    main()
